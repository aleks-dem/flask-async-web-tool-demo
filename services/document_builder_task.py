import os
import subprocess
import time
import logging
from datetime import datetime

import docx
from celery_app import celery

from utils.redis_state_repository import get_process_state, set_process_state
from utils.observability import set_request_context, reset_request_context, observe_task
import utils.globals as g


logger = logging.getLogger(__name__)


def is_highlighted(run):
    highlight = run.font.highlight_color
    return bool(highlight and str(highlight).lower() in ['yellow (7)', 'yellow', '7'])


def replace_text_in_runs(runs, new_text):
    if not runs:
        return
    for run in runs[1:]:
        run.text = ''
    runs[0].text = new_text
    runs[0].font.highlight_color = None


def replace_highlighted_text(paragraph, replacements):
    combined_text = ''
    highlighted_runs = []

    for run in paragraph.runs:
        if is_highlighted(run):
            combined_text += run.text
            highlighted_runs.append(run)
        else:
            if combined_text and combined_text in replacements:
                replace_text_in_runs(highlighted_runs, replacements[combined_text])
            combined_text = ''
            highlighted_runs = []

    if combined_text and combined_text in replacements:
        replace_text_in_runs(highlighted_runs, replacements[combined_text])


def update_process_state(
    user_id,
    function_key,
    status=None,
    progress=None,
    error=None,
    parameters=None,
    message=None,
    file_url=None,
    timestamp=None,
):
    state = get_process_state(user_id, function_key) or {}
    if status is not None:
        state['status'] = status
    if progress is not None:
        state['progress'] = progress
    if error is not None:
        state['error'] = error
    if parameters is not None:
        state['parameters'] = parameters
    if message is not None:
        state['message'] = message
    if file_url is not None:
        state['file_url'] = file_url
    if timestamp is not None:
        state['timestamp'] = timestamp
    set_process_state(user_id, function_key, state)


def create_default_document(output_path, parameters):
    document = docx.Document()
    document.add_heading('Automated Document', level=1)
    document.add_paragraph(
        'This file was generated in demo mode. Replace mock data retrieval with your production logic.'
    )

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Value'

    for item in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['name'])
        row_cells[1].text = str(item['value'])

    document.save(output_path)


def convert_docx_to_pdf(docx_path):
    subprocess.run(
        ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', g.temp_dir, docx_path],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return os.path.splitext(docx_path)[0] + '.pdf'


def generate_document_from_template(template_path, output_path, replacements):
    document = docx.Document(template_path)

    for paragraph in document.paragraphs:
        replace_highlighted_text(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_highlighted_text(paragraph, replacements)

    document.save(output_path)


@celery.task
def build_document_process(
    user_id,
    function_key,
    workspace,
    template,
    templates_path,
    reference_id,
    file_format,
    approval_code,
    network_reference,
    retrieval_reference,
    external_order_id,
):
    task_started_at = time.perf_counter()
    task_id = getattr(build_document_process.request, 'id', None)
    ctx_tokens = set_request_context(
        request_id=f'celery-{task_id}' if task_id else 'celery-unknown',
        user_id=str(user_id),
    )
    observe_task('document_builder', 'started')
    logger.info(
        'task_started',
        extra={
            'event': 'task_started',
            'function_key': function_key,
            'task_id': task_id or '-',
        },
    )

    try:
        update_process_state(user_id, function_key, status='processing', progress=10)
        time.sleep(0.4)

        update_process_state(user_id, function_key, status='processing', progress=35)
        today = datetime.utcnow().strftime('%Y-%m-%d')
        currency = 'USD' if workspace == 'workspace_a' else 'EUR'
        amount_value = f'{(reference_id % 9000) / 100 + 10:.2f}'

        replacements = {
            '{{CURRENT_DATE}}': today,
            '{{REFERENCE_ID}}': str(reference_id),
            '{{WORKSPACE}}': workspace,
            '{{AMOUNT}}': amount_value,
            '{{CURRENCY}}': currency,
            '{{APPROVAL_CODE}}': approval_code or 'N/A',
            '{{NETWORK_REFERENCE}}': network_reference or 'N/A',
            '{{RETRIEVAL_REFERENCE}}': retrieval_reference or 'N/A',
            '{{EXTERNAL_ORDER_ID}}': external_order_id or 'N/A',
            '{{STATUS}}': 'APPROVED',
        }

        parameters = [
            {'name': 'Reference ID', 'value': reference_id},
            {'name': 'Workspace', 'value': workspace},
            {'name': 'Record Date', 'value': today},
            {'name': 'Amount', 'value': f'{amount_value} {currency}'},
            {'name': 'Approval Code', 'value': approval_code or '-'},
            {'name': 'Network Reference', 'value': network_reference or '-'},
            {'name': 'Retrieval Reference', 'value': retrieval_reference or '-'},
            {'name': 'External Order ID', 'value': external_order_id or '-'},
        ]

        update_process_state(user_id, function_key, status='processing', progress=70)

        docx_filename = f'{reference_id}_{int(time.time())}.docx'
        docx_path = os.path.join(g.temp_dir, docx_filename)
        template_path = os.path.join(templates_path, f'{template}.docx')

        if os.path.exists(template_path):
            generate_document_from_template(template_path, docx_path, replacements)
        else:
            create_default_document(docx_path, parameters)

        final_path = docx_path
        conversion_note = None

        if file_format == 'pdf':
            try:
                final_path = convert_docx_to_pdf(docx_path)
                os.remove(docx_path)
            except Exception:
                final_path = docx_path
                conversion_note = 'PDF conversion is unavailable in this environment. DOCX was generated instead.'

        update_process_state(user_id, function_key, status='processing', progress=100)
        time.sleep(0.5)

        done_message = 'Document generated successfully.'
        if conversion_note:
            done_message = f'{done_message} {conversion_note}'

        update_process_state(
            user_id,
            function_key,
            status='completed',
            progress=100,
            parameters=parameters,
            message=done_message,
            file_url=f"download/{function_key}/{os.path.basename(final_path)}",
            timestamp=time.time(),
        )
        duration = max(0.0, time.perf_counter() - task_started_at)
        observe_task('document_builder', 'completed', duration_seconds=duration)
        logger.info(
            'task_completed',
            extra={
                'event': 'task_completed',
                'function_key': function_key,
                'task_id': task_id or '-',
                'duration_ms': round(duration * 1000, 2),
            },
        )
    except Exception as exc:
        duration = max(0.0, time.perf_counter() - task_started_at)
        observe_task('document_builder', 'error', duration_seconds=duration)
        logger.exception(
            'task_failed',
            extra={
                'event': 'task_failed',
                'function_key': function_key,
                'task_id': task_id or '-',
                'duration_ms': round(duration * 1000, 2),
            },
        )
        update_process_state(
            user_id,
            function_key,
            status='error',
            progress=0,
            error=str(exc),
            timestamp=time.time(),
        )
    finally:
        reset_request_context(ctx_tokens)
